"""AI Assignment Maker — turn free text / PDF / DOCX / image into a gig draft.

Admin reviews the pre-filled draft on the frontend and creates the gig via the
normal POST /api/gigs endpoint, then optionally blasts it.
"""
import base64
import io
import json
import os
import uuid
from datetime import datetime, timezone
from typing import Optional

from dotenv import load_dotenv
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile

from config import logger
from auth_deps import require_admin

load_dotenv()
router = APIRouter()

MODELS = {
    "gpt-5.5": ("openai", "gpt-5.5"),
    "claude-sonnet-4-6": ("anthropic", "claude-sonnet-4-6"),
}
MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_DOC_CHARS = 20000

SYSTEM_PROMPT = """You extract structured job assignment data for HCOB Network, a cleaning/labor/driver dispatch platform.
Today is {today}. Resolve relative dates ("this Friday", "tomorrow", "next Tuesday 9am") to concrete future dates.

Return ONLY a raw JSON object — no markdown fences, no commentary — with exactly these keys:
- "title": short job title, e.g. "Office Deep Clean — Downtown"
- "description": clear worker-facing description: the work, requirements, what to bring, special instructions
- "category": one of "cleaning", "labor", "driver"
- "location": short PUBLIC area preview like "Downtown · 21201" (neighborhood/street + zip, NOT the full address)
- "address_line": full street address if given, else null
- "scheduled_local": start date+time as "YYYY-MM-DDTHH:mm" (24h, local time), else null
- "pay_rate": number only (no $ sign), else null
- "pay_type": "hourly" or "flat" (default "hourly")
- "slots": integer count of workers needed (default 1)
- "duration_hours": number or null
- "contact_phone": string or null
- "ai_notes": 1-2 short sentences about anything ambiguous or assumed
- "missing_fields": array of key fields you could NOT find, chosen from ["title","description","location","scheduled_local","pay_rate"]

Rules: never invent addresses or pay rates — use null when absent. Keep description under 900 characters."""


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages[:15])


def _extract_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_json_obj(raw: str) -> Optional[dict]:
    s = str(raw).strip()
    start = s.find("{")
    end = s.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        return json.loads(s[start:end + 1])
    except json.JSONDecodeError:
        return None


def _sanitize_draft(d: dict) -> dict:
    def txt(k, limit=2000):
        v = d.get(k)
        return str(v).strip()[:limit] if v not in (None, "", "null") else None

    def num(k):
        try:
            v = float(d.get(k))
            return v if v >= 0 else None
        except (TypeError, ValueError):
            return None

    category = txt("category") or "cleaning"
    if category not in ("cleaning", "labor", "driver"):
        category = "cleaning"
    pay_type = txt("pay_type") or "hourly"
    if pay_type not in ("hourly", "flat"):
        pay_type = "hourly"
    try:
        slots = max(1, int(d.get("slots") or 1))
    except (TypeError, ValueError):
        slots = 1
    missing = d.get("missing_fields")
    if not isinstance(missing, list):
        missing = []
    return {
        "title": txt("title", 140),
        "description": txt("description", 2000),
        "category": category,
        "location": txt("location", 120),
        "address_line": txt("address_line", 240),
        "scheduled_local": txt("scheduled_local", 16),
        "pay_rate": num("pay_rate"),
        "pay_type": pay_type,
        "slots": slots,
        "duration_hours": num("duration_hours"),
        "contact_phone": txt("contact_phone", 30),
        "ai_notes": txt("ai_notes", 400),
        "missing_fields": [str(m) for m in missing][:6],
    }


@router.post("/admin/ai-assignments/parse")
async def parse_assignment(
    text: Optional[str] = Form(None),
    model: str = Form("gpt-5.5"),
    file: Optional[UploadFile] = File(None),
    admin: dict = Depends(require_admin),
):
    if model not in MODELS:
        raise HTTPException(400, "Unknown model")
    if not (text and text.strip()) and not file:
        raise HTTPException(400, "Type some details or upload a document")

    doc_text = ""
    image_b64 = None
    filename = None
    if file:
        data = await file.read()
        if len(data) > MAX_FILE_BYTES:
            raise HTTPException(400, "File too large (max 15MB)")
        filename = file.filename or "upload"
        ct = (file.content_type or "").lower()
        low = filename.lower()
        try:
            if "pdf" in ct or low.endswith(".pdf"):
                doc_text = _extract_pdf(data)
            elif low.endswith(".docx") or "wordprocessingml" in ct:
                doc_text = _extract_docx(data)
            elif ct.startswith("image/"):
                image_b64 = base64.b64encode(data).decode()
            else:
                raise HTTPException(400, "Unsupported file — use PDF, Word (.docx), or an image")
        except HTTPException:
            raise
        except Exception as e:
            logger.warning(f"AI assignment file extraction failed ({filename}): {e}")
            raise HTTPException(400, "Could not read that file — try a different format")

    parts = []
    if text and text.strip():
        parts.append(f"ADMIN NOTES:\n{text.strip()}")
    if doc_text.strip():
        parts.append(f"DOCUMENT CONTENT ({filename}):\n{doc_text[:MAX_DOC_CHARS]}")
    if image_b64:
        parts.append("A photo/screenshot of the work order is attached — read every detail from it.")
    user_text = "\n\n".join(parts) or "Extract the assignment from the attached image."

    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        raise HTTPException(503, "AI is not configured on this environment")

    from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent

    provider, model_name = MODELS[model]
    today = datetime.now(timezone.utc).strftime("%A, %B %d, %Y")
    chat = LlmChat(
        api_key=api_key,
        session_id=f"ai-gig::{admin['user_id']}::{uuid.uuid4().hex[:8]}",
        system_message=SYSTEM_PROMPT.format(today=today),
    ).with_model(provider, model_name)

    if image_b64:
        msg = UserMessage(text=user_text, file_contents=[ImageContent(image_b64)])
    else:
        msg = UserMessage(text=user_text)

    # Blocking send is intentional — we must parse the complete JSON draft.
    try:
        raw = await chat.send_message(msg)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"AI assignment parse LLM call failed: {e}")
        raise HTTPException(502, "Could not reach the AI right now — try again in a minute")

    parsed = _extract_json_obj(raw)
    if parsed is None:
        logger.warning(f"AI assignment parse returned non-JSON: {str(raw)[:300]}")
        raise HTTPException(502, "The AI returned an unreadable draft — try again")

    return {"draft": _sanitize_draft(parsed), "model_used": model}
